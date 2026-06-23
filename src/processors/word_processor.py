import os
import docx

from src.models.document import Document
from src.utils.text_splitter import RecursiveCharacterTextSplitter
from .base.document_processor import DocumentProcessor, DEFAULT_CHUNK_SIZE, DEFAULT_CHUNK_OVERLAP

try:
    import textract
except ImportError:
    textract = None


class WordProcessor(DocumentProcessor):
    """Word document (.doc / .docx) processor."""

    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=DEFAULT_CHUNK_SIZE,
            chunk_overlap=DEFAULT_CHUNK_OVERLAP,
        )

    def process(self, file_path: str) -> list[Document]:
        if file_path.lower().endswith(".docx"):
            text = self._extract_docx(file_path)
        else:
            text = self._extract_doc(file_path)

        doc = Document(page_content=text, metadata={"source": file_path})
        return self.text_splitter.split_documents([doc])

    def _extract_docx(self, file_path: str) -> str:
        doc = docx.Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)

    def _extract_doc(self, file_path: str) -> str:
        if textract is None:
            raise ImportError(
                "textract is required to process .doc files. "
                "Install it with: pip install textract"
            )
        return textract.process(file_path).decode("utf-8")
